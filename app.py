import re
import textwrap
from collections import Counter

import streamlit as st
import plotly.graph_objects as go

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


st.set_page_config(
    page_title="Smart Resume Analyzer",
    page_icon="",
    layout="wide",
    initial_sidebar_state="collapsed"
)


def extract_pdf_text(uploaded_file):
    if PdfReader is None:
        return ""

    try:
        reader = PdfReader(uploaded_file)
        pages = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

        return "\n".join(pages)
    except Exception:
        return ""


def extract_docx_text(uploaded_file):
    if Document is None:
        return ""

    try:
        document = Document(uploaded_file)

        paragraphs = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        return "\n".join(paragraphs)
    except Exception:
        return ""


def extract_resume_text(uploaded_file):
    if uploaded_file is None:
        return ""

    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        return extract_pdf_text(uploaded_file)

    if file_name.endswith(".docx"):
        return extract_docx_text(uploaded_file)

    return ""


def clean_text(text):
    text = text.lower()
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_skill(skill):
    return re.sub(r"[^a-z0-9+#.\- ]", "", skill.lower()).strip()


SKILLS = [
    "python",
    "java",
    "javascript",
    "typescript",
    "c",
    "c++",
    "c#",
    "react",
    "node.js",
    "nodejs",
    "fastapi",
    "flask",
    "django",
    "streamlit",
    "html",
    "css",
    "sql",
    "mysql",
    "sqlite",
    "postgresql",
    "mongodb",
    "sqlalchemy",
    "rest api",
    "rest apis",
    "api",
    "jwt",
    "git",
    "github",
    "docker",
    "aws",
    "azure",
    "gcp",
    "machine learning",
    "artificial intelligence",
    "ai",
    "nlp",
    "natural language processing",
    "pandas",
    "numpy",
    "scikit-learn",
    "tensorflow",
    "pytorch",
    "data structures",
    "algorithms",
    "oop",
    "object oriented programming",
    "testing",
    "unit testing",
    "postman",
    "vercel",
    "render"
]


def extract_skills(text):
    text_lower = clean_text(text)
    found = []

    for skill in SKILLS:
        skill_lower = skill.lower()

        if skill_lower in text_lower:
            if skill_lower == "ai":
                if re.search(r"\bai\b", text_lower):
                    found.append(skill)
            else:
                found.append(skill)

    unique = []
    for skill in found:
        if skill not in unique:
            unique.append(skill)

    return unique


def calculate_skill_match(resume_skills, job_skills):
    if not job_skills:
        return 0

    matched = set(resume_skills).intersection(set(job_skills))
    return round((len(matched) / len(set(job_skills))) * 100)


def find_matched_skills(resume_skills, job_skills):
    return [
        skill for skill in job_skills
        if skill in resume_skills
    ]


def find_missing_skills(resume_skills, job_skills):
    return [
        skill for skill in job_skills
        if skill not in resume_skills
    ]


def contains_section(text, section_names):
    text_lower = clean_text(text)

    for section in section_names:
        if section in text_lower:
            return True

    return False


def calculate_resume_quality(text):
    if not text.strip():
        return 0

    score = 0

    sections = {
        "contact": [
            "email",
            "@"
        ],
        "education": [
            "education",
            "b.tech",
            "btech",
            "bachelor",
            "degree"
        ],
        "experience": [
            "experience",
            "internship",
            "work experience"
        ],
        "projects": [
            "projects",
            "project"
        ],
        "skills": [
            "skills",
            "technical skills"
        ],
        "certifications": [
            "certification",
            "certifications",
            "certificate"
        ],
        "linkedin": [
            "linkedin"
        ],
        "github": [
            "github"
        ]
    }

    for values in sections.values():
        if contains_section(text, values):
            score += 10

    if len(text.split()) >= 150:
        score += 10

    if len(text.split()) <= 700:
        score += 10

    return min(score, 100)


def count_bullet_points(text):
    patterns = [
        r"(?m)^\s*[•●▪◦*-]\s+",
        r"(?m)^\s*\d+[.)]\s+"
    ]

    total = 0

    for pattern in patterns:
        total += len(re.findall(pattern, text))

    return total


def calculate_section_score(text, section_names):
    return 100 if contains_section(text, section_names) else 0


def calculate_project_strength(text):
    project_present = contains_section(
        text,
        ["projects", "project"]
    )

    if not project_present:
        return 0

    project_keywords = [
        "developed",
        "built",
        "implemented",
        "created",
        "designed",
        "deployed",
        "using",
        "api",
        "database",
        "github"
    ]

    text_lower = clean_text(text)

    matches = sum(
        1 for keyword in project_keywords
        if keyword in text_lower
    )

    return min(50 + matches * 5, 100)


def calculate_experience_strength(text):
    if contains_section(
        text,
        ["experience", "internship", "work experience"]
    ):
        return 80

    return 0


def calculate_education_strength(text):
    if contains_section(
        text,
        ["education", "b.tech", "btech", "bachelor", "degree"]
    ):
        return 65

    return 0


def calculate_ats_score(skill_match, quality_score, project_score, education_score):
    score = (
        skill_match * 0.45
        + quality_score * 0.20
        + project_score * 0.20
        + education_score * 0.15
    )

    return round(min(score, 100))


def extract_keywords(text):
    text_lower = clean_text(text)

    keywords = []

    for skill in SKILLS:
        if skill.lower() in text_lower:
            keywords.append(skill)

    words = re.findall(r"\b[a-zA-Z][a-zA-Z+#.-]{2,}\b", text_lower)

    stop_words = {
        "the",
        "and",
        "for",
        "with",
        "that",
        "this",
        "are",
        "you",
        "your",
        "from",
        "will",
        "our",
        "have",
        "has",
        "using",
        "into",
        "about",
        "work",
        "working",
        "candidate",
        "candidates",
        "required",
        "requirements",
        "responsibilities"
    }

    frequency = Counter(
        word for word in words
        if word not in stop_words
    )

    for word, count in frequency.most_common(30):
        if count >= 2 and word not in keywords:
            keywords.append(word)

    return keywords[:25]


def detect_job_level(job_text):
    text = clean_text(job_text)

    if any(
        phrase in text
        for phrase in [
            "senior developer",
            "senior software",
            "lead developer",
            "tech lead",
            "5+ years",
            "6+ years",
            "7+ years"
        ]
    ):
        return "Senior"

    if any(
        phrase in text
        for phrase in [
            "mid-level",
            "mid level",
            "3+ years",
            "4+ years"
        ]
    ):
        return "Mid-Level"

    if any(
        phrase in text
        for phrase in [
            "intern",
            "internship",
            "fresher",
            "entry level",
            "entry-level",
            "graduate",
            "0-1 years"
        ]
    ):
        return "Entry Level"

    return "Not specified"


def detect_experience(job_text):
    matches = re.findall(
        r"(\d+)\+?\s*(?:years?|yrs?)",
        job_text.lower()
    )

    if matches:
        return f"{matches[0]}+ years"

    return "Not specified"


def extract_responsibilities(job_text):
    lines = job_text.splitlines()

    responsibilities = []

    keywords = [
        "develop",
        "build",
        "create",
        "implement",
        "design",
        "maintain",
        "test",
        "debug",
        "integrate",
        "work with",
        "collaborate"
    ]

    for line in lines:
        cleaned = line.strip()

        if not cleaned:
            continue

        lower = cleaned.lower()

        if any(keyword in lower for keyword in keywords):
            cleaned = re.sub(
                r"^[•●▪◦*-]\s*",
                "",
                cleaned
            )

            responsibilities.append(cleaned)

    return responsibilities[:10]


def extract_preferred_skills(job_text):
    text = clean_text(job_text)

    preferred = []

    markers = [
        "preferred",
        "nice to have",
        "plus",
        "bonus"
    ]

    if any(marker in text for marker in markers):
        for skill in SKILLS:
            if skill.lower() in text:
                preferred.append(skill)

    return list(dict.fromkeys(preferred))


def build_recommendations(missing_skills):
    recommendations = []

    for skill in missing_skills[:8]:
        recommendations.append(
            f"Learn and practice {skill} through a small practical project."
        )

    if not recommendations:
        recommendations.append(
            "Your resume covers the main detected job requirements."
        )

    return recommendations


def build_improvements(
    resume_text,
    missing_skills,
    job_text
):
    improvements = []

    if missing_skills:
        improvements.append(
            "Add relevant skills from the job description only if you genuinely have experience with them."
        )

    if not contains_section(resume_text, ["experience", "internship"]):
        improvements.append(
            "Add relevant internship, academic, freelance or practical experience."
        )

    if not contains_section(resume_text, ["projects", "project"]):
        improvements.append(
            "Add 2 to 3 strong technical projects with measurable results."
        )

    if not contains_section(resume_text, ["certifications", "certification"]):
        improvements.append(
            "Include relevant technical certifications if available."
        )

    improvements.append(
        "Tailor project and experience descriptions to the important keywords in the job description."
    )

    improvements.append(
        "Use concise action-oriented bullet points and include measurable outcomes where possible."
    )

    return list(dict.fromkeys(improvements))


def create_gauge(score):
    figure = go.Figure(
        go.Indicator(
            mode="gauge+number",
            value=score,
            title={"text": "Resume Match Score"},
            gauge={
                "axis": {
                    "range": [0, 100],
                    "tickwidth": 1
                },
                "bar": {
                    "color": "white"
                },
                "steps": [
                    {
                        "range": [0, 40],
                        "color": "#ff4b5c"
                    },
                    {
                        "range": [40, 60],
                        "color": "#f9c846"
                    },
                    {
                        "range": [60, 80],
                        "color": "#62f39b"
                    },
                    {
                        "range": [80, 100],
                        "color": "#00c853"
                    }
                ]
            }
        )
    )

    figure.update_layout(
        height=360,
        margin=dict(
            l=20,
            r=20,
            t=60,
            b=20
        ),
        paper_bgcolor="#1f1f32",
        font=dict(color="white")
    )

    return figure


st.markdown(
    """
    <style>
    .main {
        padding-top: 1rem;
        padding-bottom: 3rem;
    }

    .block-container {
        max-width: 1400px;
        padding-left: 2rem;
        padding-right: 2rem;
        padding-bottom: 4rem;
    }

    .section-title {
        font-size: 1.35rem;
        font-weight: 700;
        margin-top: 2rem;
        margin-bottom: 0.9rem;
        color: #f1f5f9;
    }

    .custom-card {
        width: 100%;
        min-width: 0;
        box-sizing: border-box;
        border: 1px solid #303846;
        border-radius: 10px;
        padding: 16px;
        background: #151a22;
        color: #f1f5f9;
        overflow: hidden;
    }

    .custom-card-title {
        font-size: 1rem;
        font-weight: 700;
        margin-bottom: 12px;
        color: #f8fafc;
    }

    .custom-card-text {
        font-size: 0.9rem;
        line-height: 1.55;
        color: #e2e8f0;
        overflow-wrap: anywhere;
        word-break: normal;
    }

    .summary-grid {
        display: grid;
        grid-template-columns: repeat(4, minmax(0, 1fr));
        gap: 14px;
        width: 100%;
        align-items: start;
    }

    .summary-grid > * {
        min-width: 0;
    }

    .summary-card {
        min-height: 0;
        height: auto;
    }

    .summary-list,
    .project-list,
    .skill-list,
    .job-list,
    .recommendation-list {
        margin: 0;
        padding-left: 18px;
    }

    .summary-list li,
    .project-list li,
    .skill-list li,
    .job-list li,
    .recommendation-list li {
        margin-bottom: 8px;
        line-height: 1.5;
        color: #e2e8f0;
        overflow-wrap: anywhere;
    }

    .summary-list li:last-child,
    .project-list li:last-child,
    .skill-list li:last-child,
    .job-list li:last-child,
    .recommendation-list li:last-child {
        margin-bottom: 0;
    }

    .metrics-grid {
        display: grid;
        grid-template-columns: repeat(5, minmax(0, 1fr));
        gap: 14px;
        width: 100%;
        align-items: stretch;
    }

    .metrics-grid.three {
        grid-template-columns: repeat(3, minmax(0, 1fr));
    }

    .metric-card {
        min-width: 0;
        min-height: 96px;
        height: auto;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
    }

    .metric-label {
        font-size: 0.82rem;
        color: #cbd5e1;
        margin-bottom: 14px;
    }

    .metric-value {
        font-size: 1.45rem;
        line-height: 1.1;
        font-weight: 700;
        color: #4f73ff;
    }

    .two-column-grid {
        display: grid;
        grid-template-columns: repeat(2, minmax(0, 1fr));
        gap: 14px;
        width: 100%;
        align-items: start;
    }

    .two-column-grid > * {
        min-width: 0;
    }

    .three-column-grid {
        display: grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap: 14px;
        width: 100%;
        align-items: start;
    }

    .three-column-grid > * {
        min-width: 0;
    }

    .project-grid {
        display: grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap: 14px;
        width: 100%;
        align-items: stretch;
    }

    .project-card {
        height: 100%;
        min-height: 0;
        display: flex;
        flex-direction: column;
    }

    .skill-gap-metrics {
        display: grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap: 14px;
        width: 100%;
        margin-bottom: 14px;
    }

    .skill-gap-metrics > * {
        min-width: 0;
    }

    .skill-gap-columns {
        display: grid;
        grid-template-columns: repeat(2, minmax(0, 1fr));
        gap: 14px;
        width: 100%;
        align-items: start;
        margin-bottom: 14px;
    }

    .skill-gap-columns > * {
        min-width: 0;
    }

    .skill-gap-card {
        min-height: 0;
        height: auto;
    }

    .priority-grid {
        display: grid;
        grid-template-columns: repeat(2, minmax(0, 1fr));
        gap: 14px;
        width: 100%;
        align-items: start;
    }

    .priority-grid > * {
        min-width: 0;
    }

    .message-card {
        margin-top: 12px;
        padding: 13px 16px;
        border-radius: 8px;
        border: 1px solid #1f6f4a;
        background: #123b2b;
        color: #6ee7a8;
        line-height: 1.5;
    }

    .status-grid {
        display: grid;
        grid-template-columns: repeat(4, minmax(0, 1fr));
        gap: 14px;
        width: 100%;
    }

    .status-item {
        min-width: 0;
        border: 1px solid #303846;
        border-radius: 10px;
        padding: 13px 15px;
        background: #151a22;
        color: #e2e8f0;
        font-size: 0.88rem;
        overflow-wrap: anywhere;
    }

    .status-present {
        border-color: #1f6f4a;
    }

    .status-missing {
        border-color: #60433a;
    }

    .ai-section {
        margin-top: 32px;
    }

    .ai-card {
        padding: 20px;
    }

    .final-section {
        margin-top: 30px;
    }

    .empty-text {
        color: #94a3b8;
        font-size: 0.9rem;
    }

    .inline-value {
        color: #38bdf8;
        font-weight: 600;
    }

    @media (max-width: 1100px) {
        .summary-grid {
            grid-template-columns: repeat(2, minmax(0, 1fr));
        }

        .project-grid {
            grid-template-columns: repeat(2, minmax(0, 1fr));
        }

        .metrics-grid {
            grid-template-columns: repeat(3, minmax(0, 1fr));
        }

        .status-grid {
            grid-template-columns: repeat(2, minmax(0, 1fr));
        }
    }

    @media (max-width: 800px) {
        .summary-grid,
        .project-grid,
        .two-column-grid,
        .three-column-grid,
        .skill-gap-metrics,
        .skill-gap-columns,
        .priority-grid {
            grid-template-columns: 1fr;
        }

        .metrics-grid,
        .status-grid {
            grid-template-columns: repeat(2, minmax(0, 1fr));
        }
    }

    @media (max-width: 520px) {
        .block-container {
            padding-left: 1rem;
            padding-right: 1rem;
        }

        .metrics-grid,
        .status-grid {
            grid-template-columns: 1fr;
        }
    }
    </style>
    """,
    unsafe_allow_html=True
)


def html_escape(value):
    value = str(value)
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#39;")
    )


def clean_display_lines(lines, limit=6):
    cleaned = []
    seen = set()

    for line in lines:
        value = re.sub(r"^[•●▪◦*-]\s*", "", str(line).strip())
        value = re.sub(r"^\d+[.)]\s*", "", value)
        value = re.sub(r"\s+", " ", value).strip()

        if not value:
            continue

        key = value.lower()
        if key in seen:
            continue

        seen.add(key)
        cleaned.append(value)

        if len(cleaned) >= limit:
            break

    return cleaned


def extract_resume_sections(resume_text):
    section_aliases = {
        "education": [
            "education",
            "academic background",
            "academic qualifications"
        ],
        "experience": [
            "experience",
            "work experience",
            "professional experience",
            "internship",
            "internships"
        ],
        "certifications": [
            "certifications",
            "certification",
            "certificates",
            "certificate"
        ],
        "projects": [
            "projects",
            "project",
            "personal projects",
            "academic projects"
        ]
    }

    normalized_aliases = {
        section: {alias.lower() for alias in aliases}
        for section, aliases in section_aliases.items()
    }

    sections = {
        "education": [],
        "experience": [],
        "certifications": [],
        "projects": []
    }

    current_section = None

    for raw_line in resume_text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line:
            continue

        normalized = re.sub(r"^[•●▪◦*-]\s*", "", line)
        normalized = re.sub(r"^\d+[.)]\s*", "", normalized)
        normalized = normalized.strip(" :-|_").lower()

        detected_section = None
        for section, aliases in normalized_aliases.items():
            if normalized in aliases:
                detected_section = section
                break

        if detected_section:
            current_section = detected_section
            continue

        if current_section:
            sections[current_section].append(line)

    return sections


def render_list(items, empty_message="No information found.", css_class="summary-list"):
    cleaned = clean_display_lines(items)

    if not cleaned:
        return f'<div class="empty-text">{html_escape(empty_message)}</div>'

    list_items = "".join(
        f"<li>{html_escape(item)}</li>"
        for item in cleaned
    )

    return f'<ul class="{css_class}">{list_items}</ul>'


def render_card(title, content, extra_class=""):
    return textwrap.dedent(f"""
        <div class="custom-card {extra_class}">
            <div class="custom-card-title">{html_escape(title)}</div>
            {content}
        </div>
    """).strip()


def render_metric(label, value):
    return textwrap.dedent(f"""
        <div class="custom-card metric-card">
            <div class="metric-label">{html_escape(label)}</div>
            <div class="metric-value">{html_escape(value)}</div>
        </div>
    """).strip()


def render_status(label, present):
    css_class = "status-present" if present else "status-missing"
    state = "Present" if present else "Missing"
    return textwrap.dedent(f"""
        <div class="status-item {css_class}">
            <strong>{html_escape(label)}</strong><br>
            <span>{state}</span>
        </div>
    """).strip()


if "analysis_complete" not in st.session_state:
    st.session_state.analysis_complete = False

if "resume_text" not in st.session_state:
    st.session_state.resume_text = ""

if "job_text" not in st.session_state:
    st.session_state.job_text = ""


st.title("Smart Resume Analyzer")

uploaded_file = st.file_uploader(
    "Upload Your Resume",
    type=["pdf", "docx"],
    help="Upload your resume in PDF or DOCX format."
)

if uploaded_file is not None:
    resume_text = extract_resume_text(uploaded_file)

    if resume_text.strip():
        st.session_state.resume_text = resume_text
        st.success(
            f"Resume loaded successfully: {uploaded_file.name}"
        )
    else:
        st.error(
            "Unable to extract text from this resume. Please upload a readable PDF or DOCX file."
        )


st.markdown(
    '<div class="section-title">Job Description</div>',
    unsafe_allow_html=True
)

job_text = st.text_area(
    "Paste the complete Job Description here",
    value=st.session_state.job_text,
    height=220,
    label_visibility="visible"
)

st.session_state.job_text = job_text


if st.button(
    "Analyze Resume",
    type="primary",
    use_container_width=True
):
    if not st.session_state.resume_text.strip():
        st.error("Please upload a resume first.")
        st.stop()

    if not job_text.strip():
        st.error("Please paste a Job Description.")
        st.stop()

    st.session_state.analysis_complete = True


if st.session_state.analysis_complete:
    resume_text = st.session_state.resume_text
    job_text = st.session_state.job_text

    resume_skills = extract_skills(resume_text)
    job_skills = extract_skills(job_text)

    matched_skills = find_matched_skills(resume_skills, job_skills)
    missing_skills = find_missing_skills(resume_skills, job_skills)
    skill_match = calculate_skill_match(resume_skills, job_skills)
    quality_score = calculate_resume_quality(resume_text)
    education_score = calculate_education_strength(resume_text)
    experience_score = calculate_experience_strength(resume_text)
    project_score = calculate_project_strength(resume_text)

    ats_score = calculate_ats_score(
        skill_match,
        quality_score,
        project_score,
        education_score
    )

    job_keywords = extract_keywords(job_text)

    matched_job_keywords = [
        keyword
        for keyword in job_keywords
        if keyword in resume_skills
        or keyword in clean_text(resume_text)
    ]

    missing_job_keywords = [
        keyword
        for keyword in job_keywords
        if keyword not in matched_job_keywords
    ]

    keyword_match = (
        round(len(matched_job_keywords) / len(job_keywords) * 100)
        if job_keywords
        else 0
    )

    job_level = detect_job_level(job_text)
    experience_required = detect_experience(job_text)
    responsibilities = extract_responsibilities(job_text)
    preferred_skills = extract_preferred_skills(job_text)

    resume_words = resume_text.split()
    word_count = len(resume_words)
    bullet_points = count_bullet_points(resume_text)

    improvements = build_improvements(
        resume_text,
        missing_skills,
        job_text
    )

    learning_recommendations = build_recommendations(missing_skills)

    resume_sections = extract_resume_sections(resume_text)
    education_lines = resume_sections["education"]
    experience_lines = resume_sections["experience"]
    certification_lines = resume_sections["certifications"]

    st.markdown(
        '<div class="section-title">Resume Match Score</div>',
        unsafe_allow_html=True
    )

    st.plotly_chart(
        create_gauge(ats_score),
        use_container_width=True
    )

    st.markdown(
        '<div class="section-title">ATS Resume Breakdown</div>',
        unsafe_allow_html=True
    )

    ats_metrics = [
        ("ATS Compatibility", ats_score),
        ("Skills Match", skill_match),
        ("Education Strength", education_score),
        ("Experience Strength", experience_score),
        ("Project Strength", project_score)
    ]

    st.markdown(
        '<div class="metrics-grid">' +
        "".join(
            render_metric(label, f"{value}%")
            for label, value in ats_metrics
        ) +
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-title">Job Keywords</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        render_card(
            "Detected Keywords",
            render_list(
                job_keywords,
                "No significant keywords detected.",
                "skill-list"
            )
        ),
        unsafe_allow_html=True
    )

    st.markdown(
        textwrap.dedent(f"""
            <div class="two-column-grid">
                {render_card("Matched Job Keywords", render_list(matched_job_keywords, "No matched keywords found.", "skill-list"))}
                {render_card("Missing Job Keywords", render_list(missing_job_keywords, "No major missing keywords detected.", "skill-list"))}
            </div>
        """).strip(),
        unsafe_allow_html=True
    )

    if preferred_skills:
        preferred_missing = [
            skill
            for skill in preferred_skills
            if skill not in resume_skills
        ]

        st.markdown(
            '<div class="section-title">Preferred Skills Gap</div>',
            unsafe_allow_html=True
        )

        st.markdown(
            render_card(
                "Preferred Skills You Are Missing",
                render_list(
                    preferred_missing,
                    "No preferred skill gaps detected.",
                    "skill-list"
                )
            ),
            unsafe_allow_html=True
        )

    st.markdown(
        '<div class="section-title">Resume Quality Analysis</div>',
        unsafe_allow_html=True
    )

    quality_metrics = [
        ("Resume Quality Score", quality_score),
        ("Word Count", word_count),
        ("Bullet Points", bullet_points)
    ]

    st.markdown(
        '<div class="metrics-grid three">' +
        "".join(
            render_metric(
                label,
                f"{value}%" if label == "Resume Quality Score" else str(value)
            )
            for label, value in quality_metrics
        ) +
        '</div>',
        unsafe_allow_html=True
    )

    if word_count < 150:
        quality_message = "Resume may be too short for a complete application."
    elif word_count > 800:
        quality_message = "Resume may be longer than necessary."
    else:
        quality_message = "Resume length is within a reasonable range."

    st.markdown(
        f'<div class="message-card">{html_escape(quality_message)}</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-title">Resume Strengths</div>',
        unsafe_allow_html=True
    )

    strengths = []

    if contains_section(resume_text, ["education"]):
        strengths.append("Education section is present.")

    if contains_section(resume_text, ["experience", "internship"]):
        strengths.append("Experience section is present.")

    if contains_section(resume_text, ["projects"]):
        strengths.append("Projects section is present.")

    if resume_skills:
        strengths.append("Technical skills are detected.")

    if contains_section(resume_text, ["certification", "certifications"]):
        strengths.append("Certifications are included.")

    if "@" in resume_text:
        strengths.append("Email address is detected.")

    if re.search(r"\+?\d[\d\s-]{8,}", resume_text):
        strengths.append("Phone number is detected.")

    if "linkedin" in clean_text(resume_text):
        strengths.append("LinkedIn profile is included.")

    if "github" in clean_text(resume_text):
        strengths.append("GitHub profile is included.")

    st.markdown(
        render_card(
            "Detected Strengths",
            render_list(
                strengths,
                "No specific strengths detected.",
                "summary-list"
            )
        ),
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-title">Areas to Improve</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        render_card(
            "Improvement Suggestions",
            render_list(
                improvements,
                "No additional suggestions available.",
                "recommendation-list"
            )
        ),
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-title">Resume Summary</div>',
        unsafe_allow_html=True
    )

    summary_cards = [
        render_card(
            "Education",
            render_list(
                education_lines,
                "No education found.",
                "summary-list"
            ),
            "summary-card"
        ),
        render_card(
            "Skills",
            render_list(
                resume_skills[:12],
                "No skills found.",
                "skill-list"
            ),
            "summary-card"
        ),
        render_card(
            "Certificates",
            render_list(
                certification_lines,
                "No certificates found.",
                "summary-list"
            ),
            "summary-card"
        ),
        render_card(
            "Experience",
            render_list(
                experience_lines,
                "No experience found.",
                "summary-list"
            ),
            "summary-card"
        )
    ]

    st.markdown(
        '<div class="summary-grid">' +
        "".join(summary_cards) +
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-title">Projects</div>',
        unsafe_allow_html=True
    )

    project_lines = resume_sections["projects"]

    project_data = []
    current_project = None
    current_descriptions = []

    for line in project_lines:
        line = str(line).strip()

        if not line:
            continue

        cleaned_line = re.sub(r"^[•●▪◦*-]\s*", "", line).strip()

        if not cleaned_line:
            continue

        words = cleaned_line.split()

        is_description = (
            len(words) > 7
            or cleaned_line.endswith(".")
            or cleaned_line.endswith(",")
            or cleaned_line.endswith(":")
            or cleaned_line.lower().startswith((
                "built ",
                "developed ",
                "implemented ",
                "created ",
                "designed ",
                "deployed ",
                "added ",
                "used ",
                "worked ",
                "processed ",
                "generated "
            ))
        )

        if not is_description:
            if current_project:
                project_data.append(
                    (
                        current_project,
                        current_descriptions
                    )
                )

            current_project = cleaned_line
            current_descriptions = []

        else:
            if current_project:
                current_descriptions.append(cleaned_line)

    if current_project:
        project_data.append(
            (
                current_project,
                current_descriptions
            )
        )

    project_cards = [
        render_card(
            title,
            render_list(
                descriptions,
                "No project description available.",
                "project-list"
            ),
            "project-card"
        )
        for title, descriptions in project_data
    ]

    st.markdown(
        '<div class="project-grid">' +
        "".join(project_cards) +
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="section-title">Recommended Jobs</div>',
        unsafe_allow_html=True
    )

    recommended_jobs = []

    if "python" in job_skills:
        recommended_jobs.append("Python Developer")

    if "react" in job_skills:
        recommended_jobs.append("React Developer")

    if "fastapi" in job_skills:
        recommended_jobs.append("Backend Developer")

    if "sql" in job_skills:
        recommended_jobs.append("SQL / Backend Developer")

    if not recommended_jobs:
        recommended_jobs = [
            "Software Developer",
            "Python Developer"
        ]

    recommended_jobs = list(dict.fromkeys(recommended_jobs))[:5]

    st.markdown(
        render_card(
            "Recommended Roles",
            render_list(
                recommended_jobs,
                "No recommended jobs found.",
                "job-list"
            ),
            "job-card"
        ),
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="ai-section">'
        '<div class="section-title">AI Resume Improvement Generator</div>'
        '</div>',
        unsafe_allow_html=True
    )

    st.markdown(
        render_card(
            "AI Resume Improvement Generator",
            '<div class="custom-card-text">Generate personalized suggestions based on your resume and the job description.</div>',
            "ai-card"
        ),
        unsafe_allow_html=True
    )

    if st.button(
        "Generate Improvements",
        use_container_width=True,
        key="generate_improvements_button"
    ):
        st.markdown(
            render_card(
                "Suggested Improvements",
                render_list(
                    improvements,
                    "No suggestions available.",
                    "recommendation-list"
                )
            ),
            unsafe_allow_html=True
        )

    st.markdown(
        '<div class="final-section">'
        '<div class="section-title">Final Recommendation</div>'
        '</div>',
        unsafe_allow_html=True
    )

    if ats_score >= 80:
        final_message = "Your resume is strongly aligned with this job description."
        final_class = "status-present"
    elif ats_score >= 60:
        final_message = "Your resume has a reasonable match. Improving the missing skills and tailoring the resume can strengthen it."
        final_class = ""
    else:
        final_message = "Your resume has a low match with this job description. Focus on the missing skills and tailor relevant projects and experience."
        final_class = "status-missing"

    st.markdown(
        f'<div class="status-item {final_class}" style="line-height:1.6;">{html_escape(final_message)}</div>',
        unsafe_allow_html=True
    )